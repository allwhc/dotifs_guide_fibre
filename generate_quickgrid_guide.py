from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
for _ in range(4):
    doc.add_paragraph('')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('DOTIFS QuickGrid')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0, 51, 102)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('User Guide')
r.font.size = Pt(16); r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')
a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run('System Design & Development: Vishal Jain')
r.bold = True; r.font.size = Pt(12)

c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('Project Coordinator: Deepa Modi')
r.font.size = Pt(12)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('May 2026')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ── 1. Introduction ──
doc.add_heading('1. What QuickGrid Does', level=1)
doc.add_paragraph(
    'QuickGrid is the simplest way to drive fibre placement into the DOTIFS Micro-Lens Array (MLA). '
    'It guides the operator through four clear steps, manages stage positioning, and tracks the '
    'fibre status for every hole in the 12 x 12 grid.'
)
doc.add_paragraph(
    'It is designed for one-handed routine work: open the page, click through the four steps, '
    'and start placing fibres. Advanced settings stay hidden until needed.'
)

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: QuickGrid main view]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 2. Opening QuickGrid ──
doc.add_heading('2. Opening QuickGrid', level=1)
doc.add_paragraph(
    'Open the DOTIFS application in a browser (Chrome or Edge), then click the "QuickGrid" '
    'tab at the top of the page. The page shows four step circles at the top (Connect, Set '
    'Origin, Load File, Ready). Each circle turns green as the step is completed.'
)

# ── 3. The Four Steps ──
doc.add_heading('3. The Four Steps', level=1)

# Step 1
doc.add_heading('Step 1 — Connect', level=2)
doc.add_paragraph('Click "Connect". A browser dialog asks which serial port to use:')
items = [
    'Pick the COM port that is connected to the PI stage controllers',
    'Click "Connect" in the dialog',
    'The step circle turns green when connected',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Connect step, port selection dialog]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# Step 2
doc.add_heading('Step 2 — Set Origin', level=2)
doc.add_paragraph(
    'The Origin is the physical location of hole (0,0) on the MLA — the reference point that '
    'all other holes are measured from. Getting the Origin right is the most important step.'
)
doc.add_paragraph('Procedure:')
items = [
    'Use the nudge arrows (right side panel) to move the stage to hole (0,0)',
    'Look through the microscope and align the fibre insertion needle exactly over the hole',
    'Verify by attempting to physically insert the fibre — the fibre should enter cleanly',
    'Only after the fibre enters successfully, click "Set this as Origin"',
    'A prompt appears: "Do you really want to set this as the grid origin (0,0)?"',
    'Click "Yes, Lock" — the step circle turns green',
]
for i, item in enumerate(items):
    p = doc.add_paragraph(item, style='List Number')

p = doc.add_paragraph()
r = p.add_run('Important: ')
r.bold = True; r.font.color.rgb = RGBColor(180, 0, 0)
p.add_run(
    'A small error in the origin (even 5 micrometers) will offset ALL 144 holes by the '
    'same amount. Spend extra time on this step. Physically inserting the fibre into hole '
    '(0,0) is the most reliable way to confirm alignment.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Set Origin prompt with X, Y values]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# Step 3
doc.add_heading('Step 3 — Load File', level=2)
doc.add_paragraph(
    'Click "Load File" and choose the MLA coordinate file. Two types of files are supported, '
    'and QuickGrid detects which one automatically:'
)

doc.add_heading('Pixel File (from image processing)', level=3)
doc.add_paragraph(
    'This is the original file containing the centroid pixel coordinates of each hole '
    'detected from the MLA image. When you load this:'
)
items = [
    'QuickGrid converts each pixel coordinate to physical micrometers using the formula:',
    '          distance (microns) = sqrt(dx^2 + dy^2) x 3.45',
    'The factor 3.45 is the camera scale (microns per pixel)',
    'All 144 hole positions are calculated relative to your Set Origin',
    'A new position file is automatically downloaded with a timestamp, ready for reuse',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    if 'sqrt' in item:
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0, 80, 160)
            run.bold = True

doc.add_heading('Position File (previously saved)', level=3)
doc.add_paragraph(
    'This is a file that was saved earlier from QuickGrid — already in micrometers, '
    'no conversion needed. When you load this:'
)
items = [
    'QuickGrid reads the saved Origin from the file header',
    'A prompt appears asking: "Use Same Origin" or "Use Current Origin"',
    'Choose "Use Same Origin" if the MLA is in the exact same physical mount as before',
    'Choose "Use Current Origin" if you have remounted the MLA at a new location',
    'If the file contains recorded corrections from previous work, they are restored automatically',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: File load and origin choice dialog]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# Step 4
doc.add_heading('Step 4 — Ready', level=2)
doc.add_paragraph(
    'Once all three steps above are green, you can start moving across the grid and '
    'inserting fibres. The grid shows all 144 holes as circles in their hex-staggered '
    'arrangement matching the physical MLA.'
)

# ── 4. Working With the Grid ──
doc.add_heading('4. Working With the Grid', level=1)

doc.add_heading('Click a Hole to Move', level=2)
doc.add_paragraph(
    'Click any circle on the grid to move the stage to that hole. A confirmation dialog shows '
    'exactly which commands will be sent and where the stage will end up. Click Move to proceed.'
)

doc.add_heading('Fine-Align Under the Microscope', level=2)
doc.add_paragraph(
    'After the stage arrives, use the nudge controls in the right panel to make small '
    'corrections until the fibre needle is precisely over the hole:'
)
items = [
    'Enter a step size in micrometers (try 1 for fine, 5 or 10 for coarser)',
    'Use the arrows: up, down, left, right',
    'Each click moves the stage by the chosen step size',
    'Watch the live X, Y position update in the panel',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Mark Fibre Status', level=2)
doc.add_paragraph(
    'After placing (or failing to place) a fibre, mark the hole status in the right panel:'
)
items = [
    'Yes — fibre inserted successfully (hole turns green with a tick)',
    'No — no fibre placed here (hole turns dark grey)',
    'Skip / Delete — exclude this hole from auto-scan (hole turns red with an X)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Record the Correction', level=2)
doc.add_paragraph(
    'After fine-aligning, click "Record Position" to save the small correction you made. '
    'QuickGrid stores how far you nudged from the calculated position. Next time you visit '
    'this hole (or load this file later), the stage will go directly to the corrected '
    'position without needing to re-align manually.'
)
doc.add_paragraph(
    'Recorded holes appear with an amber ring on the grid. The status bar shows how many '
    'holes have been recorded.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Point details panel with Record button and Yes/No/Skip]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 5. Saving and Reloading Your Work ──
doc.add_heading('5. Saving and Reloading Your Work', level=1)

doc.add_heading('Download Position File', level=2)
doc.add_paragraph(
    'Click "Export Recorded" at the top to download a position file containing the calculated '
    'positions plus any recorded corrections. The file name includes a timestamp.'
)
doc.add_paragraph(
    'Use this file:'
)
items = [
    'To resume work later — load it back and pick up where you left off',
    'To process a second identical MLA — load it and use the same corrections',
    'As a permanent record of where each fibre was placed',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Reload a Position File', level=2)
doc.add_paragraph(
    'Use "Load File" (Step 3) and pick the saved position file. QuickGrid asks whether to '
    'use the same origin as before or your currently locked origin. Any recorded corrections '
    'in the file are restored.'
)

# ── 6. Other Useful Buttons ──
doc.add_heading('6. Other Useful Buttons', level=1)

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = 'Button'
table.rows[0].cells[1].text = 'What it does'
for r in table.rows[0].cells[0].paragraphs[0].runs: r.bold = True
for r in table.rows[0].cells[1].paragraphs[0].runs: r.bold = True

buttons = [
    ('Speed (mm/s)', 'Sets the universal stage speed for all moves and nudges. Default 0.1 mm/s is safe for microscope work.'),
    ('Apply (speed)', 'Sends the speed to both X and Y controllers immediately.'),
    ('Acknowledge All Errors', 'Reads error codes from both stages and clears them. Use this if a stage reports an error.'),
    ('Export Recorded', 'Downloads the current position file with calculated + recorded values.'),
    ('Reset Grid', 'Clears all 144 points and their states. Origin stays locked. Confirmation required.'),
]
for i, (b, d) in enumerate(buttons):
    table.rows[i+1].cells[0].text = b
    table.rows[i+1].cells[1].text = d

# ── 7. Tips for Best Results ──
doc.add_heading('7. Tips for Best Results', level=1)
items = [
    'Always confirm the Origin by physically inserting a fibre into hole (0,0) before locking it.',
    'Use 0.1 mm/s speed during fine alignment so movements are slow and easy to follow under the microscope.',
    'Record the correction at each hole — it helps with the next MLA and prevents re-doing the alignment work.',
    'Keep the Acknowledge All Errors button in mind. If something gets stuck, it usually clears the issue.',
    'Save your work often by clicking Export Recorded. You can always reload that file to continue.',
    'If a hole has no fibre or is bad, mark it Skip / Delete so auto-scan moves past it.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── 8. Quick Reference of Hole Colors ──
doc.add_heading('8. Hole Colours on the Grid', level=1)

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = 'Colour'
table.rows[0].cells[1].text = 'Meaning'
for r in table.rows[0].cells[0].paragraphs[0].runs: r.bold = True
for r in table.rows[0].cells[1].paragraphs[0].runs: r.bold = True

colors = [
    ('Grey', 'Unvisited — not yet moved to'),
    ('Blue ring', 'Loaded — coordinates are ready, not yet visited'),
    ('Amber (glowing)', 'Current — stage is at this point right now'),
    ('Green with tick', 'Fibre placed successfully'),
    ('Red with X', 'Skipped or deleted by user'),
]
for i, (c, m) in enumerate(colors):
    table.rows[i+1].cells[0].text = c
    table.rows[i+1].cells[1].text = m

# ── Footer ──
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Prepared by: Vishal Jain')
r.bold = True
doc.add_paragraph('Project Coordinator: Deepa Modi')
doc.add_paragraph('Project: DOTIFS')
doc.add_paragraph('Date: May 2026')

output = r'c:\Users\Lab_Engineer\Downloads\Deepa maam\QuickGrid_User_Guide.docx'
doc.save(output)
print(f'Report saved: {output}')
