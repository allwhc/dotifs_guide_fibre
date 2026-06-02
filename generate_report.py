from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
for _ in range(5):
    doc.add_paragraph('')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('DOTIFS Fibre Positioning System')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0, 51, 102)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('User Guide — Grid Scan & Position Recording')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')
a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run('System Design & Development: Vishal Jain')
r.bold = True; r.font.size = Pt(12)

c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('Project Coordinator: Deepa Modi')
r.font.size = Pt(12)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('April 2026')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ── 1. Accessing the Application ──
doc.add_heading('1. Accessing the Application', level=1)
doc.add_paragraph(
    'Open a Chrome or Edge browser and navigate to:'
)
p = doc.add_paragraph()
r = p.add_run('    dotifs.netlify.app')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0, 80, 160)

doc.add_paragraph(
    'The application opens with three tabs: GCS Commands (terminal), Live Dashboard, and Grid Scan. '
    'Ensure the PI controllers are powered on and connected via USB-Serial before proceeding.'
)

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Application home page with tabs visible]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 2. Connecting to Controllers ──
doc.add_heading('2. Connecting to PI Controllers', level=1)

steps = [
    'Select the correct baud rate from the dropdown in the header (default: 115200)',
    'Click "Connect" — a browser dialog will ask you to select the serial port',
    'Choose the COM port connected to the PI controller and click "Connect"',
    'The connection indicator turns green when connected',
    'Go to the GCS Commands tab and click "Scan IDs for *IDN?" to detect all controllers in the daisy chain',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Connected state with scan results showing controller IDs]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 3. Grid Scan — Setup ──
doc.add_heading('3. Grid Scan — Initial Setup', level=1)
doc.add_paragraph('Click on the "Grid Scan" tab (password protected). After entering the password:')

steps = [
    'Set X controller ID and Y controller ID from the dropdowns (e.g., X = ID 1, Y = ID 2)',
    'Set the Axis number (default: 1)',
    'Set File units to "um (micron)" — this ensures coordinates are correctly converted to mm for the stage',
    'Check "Invert X" and/or "Invert Y" if the stage moves opposite to what you expect under the microscope',
    'Click "Load .txt positions" and select the MLA coordinate file (e.g., MLA_2_I1_3_rot_cropped.txt)',
    'The 12x12 hex grid will populate with the loaded coordinates',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Grid Scan tab with loaded coordinates and toolbar settings]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 4. Setting the Origin ──
doc.add_heading('4. Setting Your Origin (Grid 0,0 Reference)', level=1)
doc.add_paragraph(
    'Before any grid movement, you must set the origin. The origin defines where grid position (0,0) '
    'is on the physical stage. All other grid points will move relative to this position.'
)

steps = [
    'Click "Set Your Origin" button in the toolbar',
    'The origin panel opens showing live X and Y positions',
    'Use the step size presets (0.1, 1, 10, 100, 1000 um) or enter a custom value',
    'Use the arrow buttons to manually move the stage while looking through the microscope',
    'Align the fibre insertion needle precisely over the first hole (grid position 0,0)',
    'Click "Lock as Origin" — the current position is saved as the reference origin',
    'The status bar shows "Origin: X=... Y=... " with a lock icon',
    'Click "Close" to return to the grid',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Alternatively: ')
r.bold = True
p.add_run(
    'Click "FRF Both > Lock as Origin" to home both axes using the FRF command and '
    'automatically set that position as the origin.'
)

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Set Your Origin panel with jog controls and Lock button]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 5. Moving to Grid Points ──
doc.add_heading('5. Moving to Grid Points (Fibre Insertion Mode)', level=1)
doc.add_paragraph(
    'With the origin set, you can now navigate to each fibre position on the grid:'
)

steps = [
    'Click any dot on the hex grid — a confirmation dialog appears',
    'The dialog shows the exact commands that will be sent:\n'
    '    - MOV to origin (both axes)\n'
    '    - Wait for on-target\n'
    '    - MVR by the file offset (relative move from origin)',
    'Click "Move" to execute — the stage first goes to origin, then MVR to the target point',
    'Wait for "Arrived" status in the status bar',
    'Use the Fine Adjust (nudge) arrows on the right panel to make small corrections while looking through the microscope',
    'Mark the point: "Yes" (fibre inserted) or "No" (no fibre) or "Delete/Skip"',
    'Click "Next Point" to automatically move to the next unvisited point in serpentine order',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Move confirmation dialog showing command preview]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Grid with fibre/no-fibre/skipped states and nudge controls]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 6. Fine Adjust (Nudge) ──
doc.add_heading('6. Fine Adjust (Nudge Controls)', level=1)
doc.add_paragraph(
    'After the stage arrives at a grid point, use the nudge controls on the right panel '
    'to fine-tune the position while viewing through the microscope:'
)

steps = [
    'Enter step size in um (e.g., 1, 5, 65, 75, 250) or use preset buttons',
    'Enter number of steps (e.g., step = 65 um, steps = 4, total = 260 um)',
    'The "Total" line shows the exact distance that will be moved',
    'Click the arrow buttons: Up/Down for Y axis, Left/Right for X axis',
    'Each click sends one MVR command — the stage moves by the total amount',
    'The last sent command is shown at the bottom for verification',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# ── 7. Recording Positions ──
doc.add_heading('7. Recording Positions (Calibration Mode)', level=1)
doc.add_paragraph(
    'Record Mode allows you to capture the actual stage position at each grid point. '
    'This creates a calibration file that can be reused for identical MLA assemblies.'
)

doc.add_heading('7.1 How to Record', level=2)
steps = [
    'Set your origin first (Section 4)',
    'Click "Record: OFF" button in the toolbar — it toggles to "Record: ON" (amber)',
    'Manually move the stage to a fibre hole using the nudge controls',
    'Look through the microscope and align precisely',
    'Click the corresponding grid cell (the dot on the hex grid)',
    'A confirmation dialog appears: "Record position for point #N?"',
    'Click "Record" — the software reads POS? from both axes',
    'The position is stored relative to your origin (in um)',
    'The grid dot turns amber/yellow to show it has been recorded',
    'Repeat for each fibre position you want to record',
    'The status bar shows "Rec: N" — the count of recorded points',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Record Mode ON with some amber/yellow recorded points on grid]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

doc.add_heading('7.2 Exporting Recorded Positions', level=2)
steps = [
    'Click "Export Recorded .txt" button in the toolbar',
    'A .txt file downloads with only the recorded points',
    'The file includes the origin reference position in the header',
    'File format: col  row  X_um  Y_um (same as input format)',
    'Only points that were actually recorded are included (not all 144)',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

doc.add_heading('7.3 Loading a Recorded File', level=2)
steps = [
    'Click "Load .txt positions" and select a previously recorded file',
    'The software detects the origin header and shows a dialog:',
    '    "Use Same Origin" — sets origin to the recorded values (same physical setup)',
    '    "Set New Origin" — opens the origin panel to set a new reference (different setup)',
    'Recorded points appear as amber/yellow dots on the grid',
    'You can now click these points to move the stage to those recorded positions',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'Step {i+1}: {step}')

# Screenshot placeholder
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[Screenshot: Origin choice dialog when loading recorded file]')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(180, 80, 80); r.italic = True
doc.add_paragraph('')

# ── 8. Quick Reference ──
doc.add_heading('8. Quick Reference', level=1)

doc.add_heading('Grid Point Colors', level=2)
colors = [
    ('Grey', 'Unvisited — not yet moved to'),
    ('Amber (pulsing)', 'Current — stage is at this point'),
    ('Green (with tick)', 'Fibre inserted successfully'),
    ('Dark grey', 'No fibre — visited but skipped'),
    ('Red (with X)', 'Deleted/skipped — auto-scan will skip this point'),
    ('Amber (with border)', 'Recorded — position has been captured'),
]
for color, desc in colors:
    p = doc.add_paragraph()
    r = p.add_run(color + ': ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('Key Buttons', level=2)
buttons = [
    ('Set Your Origin', 'Open panel to manually position and lock the grid origin'),
    ('Go to Origin', 'Move both axes back to the locked origin position'),
    ('FRF Both XY', 'Home both axes using reference switches'),
    ('Next Point', 'Auto-move to the next unvisited point in serpentine order'),
    ('Record: ON/OFF', 'Toggle between move mode and position recording mode'),
    ('Export Recorded .txt', 'Download recorded positions as a coordinate file'),
    ('Invert X / Invert Y', 'Flip axis direction to match microscope view'),
]
for btn, desc in buttons:
    p = doc.add_paragraph()
    r = p.add_run(btn + ' — ')
    r.bold = True
    p.add_run(desc)

# ── Footer ──
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Prepared by: Vishal Jain')
r.bold = True
doc.add_paragraph('Project Coordinator: Deepa Modi')
doc.add_paragraph('Project: DOTIFS — Devasthal Optical Telescope Integral Field Spectrograph')
doc.add_paragraph('Application: dotifs.netlify.app')
doc.add_paragraph('Date: April 2026')

output = r'c:\Users\Lab_Engineer\Downloads\Deepa maam\DOTIFS_Report_v2.docx'
doc.save(output)
print(f'Report saved: {output}')
